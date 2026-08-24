from pathlib import Path

from django.conf import settings
from django.core.management.base import BaseCommand
from docx import Document

from inventario.models import Categoria, Proveedor, Producto


class Command(BaseCommand):
    help = 'Importa productos desde archivos Word de inventario'

    def leer_productos_docx(self, ruta_archivo):
        productos = []

        documento = Document(ruta_archivo)

        # Lee párrafos normales del documento.
        for parrafo in documento.paragraphs:
            texto = parrafo.text.strip()

            if not texto:
                continue

            if texto.isdigit():
                continue

            if texto.upper() in [
                'MATERIALES DE OFICINA Y PEDAGÓGICOS',
                'MATERIALES DE ASEO'
            ]:
                continue

            productos.append(texto)

        # Lee también tablas, por si el Word contiene celdas.
        for tabla in documento.tables:
            for fila in tabla.rows:
                for celda in fila.cells:
                    texto = celda.text.strip()

                    if not texto:
                        continue

                    if texto.isdigit():
                        continue

                    if texto.upper() in [
                        'MATERIALES DE OFICINA Y PEDAGÓGICOS',
                        'MATERIALES DE ASEO'
                    ]:
                        continue

                    productos.append(texto)

        # Elimina duplicados manteniendo el orden.
        productos_limpios = []
        vistos = set()

        for producto in productos:
            nombre = producto.strip()

            if nombre and nombre not in vistos:
                productos_limpios.append(nombre)
                vistos.add(nombre)

        return productos_limpios

    def handle(self, *args, **kwargs):
        carpeta_datos = Path(settings.BASE_DIR) / 'datos'

        archivo_oficina = carpeta_datos / 'oficina_pedagogicos.docx'
        archivo_aseo = carpeta_datos / 'materiales de aseo.docx'

        if not archivo_oficina.exists():
            self.stdout.write(
                self.style.ERROR(f'No se encontró el archivo: {archivo_oficina}')
            )
            return

        if not archivo_aseo.exists():
            self.stdout.write(
                self.style.ERROR(f'No se encontró el archivo: {archivo_aseo}')
            )
            return

        categoria_oficina, _ = Categoria.objects.get_or_create(
            nombre='Materiales de oficina y pedagógicos',
            defaults={
                'descripcion': 'Materiales administrativos, escolares y pedagógicos utilizados por la institución.'
            }
        )

        categoria_aseo, _ = Categoria.objects.get_or_create(
            nombre='Materiales de aseo',
            defaults={
                'descripcion': 'Insumos de limpieza, higiene y mantención de espacios institucionales.'
            }
        )

        proveedor, _ = Proveedor.objects.get_or_create(
            nombre='Proveedor Institucional',
            defaults={
                'telefono': '+56900000000',
                'correo': 'proveedor@institucional.cl',
                'direccion': 'La Pintana, Chile'
            }
        )

        productos_oficina = self.leer_productos_docx(archivo_oficina)
        productos_aseo = self.leer_productos_docx(archivo_aseo)

        total_creados = 0
        total_actualizados = 0

        for nombre_producto in productos_oficina:
            _, creado = Producto.objects.update_or_create(
                nombre=nombre_producto,
                defaults={
                    'descripcion': 'Producto importado desde listado de materiales de oficina y pedagógicos.',
                    'stock': 10,
                    'precio': 1000,
                    'categoria': categoria_oficina,
                    'proveedor': proveedor,
                }
            )

            if creado:
                total_creados += 1
            else:
                total_actualizados += 1

        for nombre_producto in productos_aseo:
            _, creado = Producto.objects.update_or_create(
                nombre=nombre_producto,
                defaults={
                    'descripcion': 'Producto importado desde listado de materiales de aseo.',
                    'stock': 10,
                    'precio': 1000,
                    'categoria': categoria_aseo,
                    'proveedor': proveedor,
                }
            )

            if creado:
                total_creados += 1
            else:
                total_actualizados += 1

        self.stdout.write(
            self.style.SUCCESS(
                f'Importación completada. Productos creados: {total_creados}. '
                f'Productos actualizados: {total_actualizados}.'
            )
        )